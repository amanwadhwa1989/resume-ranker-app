"""
utils/file_parser.py
─────────────────────
Handles extraction of plain text from uploaded resume files.
Supports PDF (via pypdf) and DOCX (via python-docx).
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import BinaryIO


# ── PDF extraction ────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes | BinaryIO) -> str:
    """
    Extract text from a PDF file.
    Falls back gracefully if a page yields no text (scanned pages).
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError("pypdf is required: pip install pypdf") from exc

    if isinstance(file_bytes, bytes):
        file_bytes = io.BytesIO(file_bytes)

    reader = PdfReader(file_bytes)
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text.strip())

    full_text = "\n\n".join(p for p in pages if p)
    return _clean_text(full_text)


# ── DOCX extraction ───────────────────────────────────────────────────────────

def extract_text_from_docx(file_bytes: bytes | BinaryIO) -> str:
    """
    Extract text from a DOCX file, preserving paragraph breaks.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("python-docx is required: pip install python-docx") from exc

    if isinstance(file_bytes, bytes):
        file_bytes = io.BytesIO(file_bytes)

    doc = Document(file_bytes)
    paragraphs: list[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Tables (skills matrices, project tables, etc.)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if row_cells:
                paragraphs.append(" | ".join(row_cells))

    full_text = "\n".join(paragraphs)
    return _clean_text(full_text)


# ── Dispatcher ────────────────────────────────────────────────────────────────

def extract_resume_text(uploaded_file) -> str:
    """
    Streamlit UploadedFile → clean plain-text string.
    Dispatches to the correct extractor based on file extension.

    Args:
        uploaded_file: A Streamlit UploadedFile object (has .name and .read()).

    Returns:
        Extracted plain text of the resume.

    Raises:
        ValueError: If the file type is unsupported.
        RuntimeError: If extraction fails.
    """
    file_bytes = uploaded_file.read()
    ext = Path(uploaded_file.name).suffix.lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext in {".docx", ".doc"}:
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'. Please upload a PDF or DOCX file."
        )

    if not text or len(text.strip()) < 50:
        raise RuntimeError(
            "Could not extract readable text from the uploaded file. "
            "If this is a scanned PDF, please use a text-based version."
        )

    return text


# ── Text cleaning ─────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """
    Normalise extracted text:
    - Collapse excessive blank lines
    - Remove non-printable characters
    - Strip trailing whitespace per line
    """
    # Remove non-printable except newlines/tabs
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)

    # Normalise line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Strip trailing whitespace per line
    lines = [line.rstrip() for line in text.split("\n")]

    # Collapse 3+ consecutive blank lines into 2
    cleaned: list[str] = []
    blank_count = 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count <= 2:
                cleaned.append(line)
        else:
            blank_count = 0
            cleaned.append(line)

    return "\n".join(cleaned).strip()


# ── Word count helper ─────────────────────────────────────────────────────────

def word_count(text: str) -> int:
    return len(text.split())
